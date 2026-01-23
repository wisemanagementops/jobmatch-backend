"""
Minimal Safe DOCX Editor

This editor ONLY adds content to an existing resume.
It does NOT:
- Delete any content
- Change any formatting
- Modify any styling
- Alter any spacing

It ONLY:
- Appends keywords to the skills section
- Optionally appends keywords to relevant bullet points
"""

import io
from typing import List
from docx import Document


def create_tailored_resume_from_template(
    original_docx_bytes: bytes,
    ai_suggestions: dict
) -> bytes:
    """
    Safely modify a resume by ONLY ADDING keywords.
    Preserves all original formatting, spacing, and content.
    """
    keywords_to_add = ai_suggestions.get('keywords', [])
    
    if not keywords_to_add:
        print("No keywords to add, returning original")
        return original_docx_bytes
    
    try:
        # Load the original document
        doc = Document(io.BytesIO(original_docx_bytes))
        
        print(f"=== SAFE DOCX EDITOR ===")
        print(f"Keywords to add: {keywords_to_add}")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        # Track what we did
        modifications = []
        
        # Step 1: Find skills section and append keywords
        skills_modified = append_to_skills_section(doc, keywords_to_add)
        if skills_modified:
            modifications.append(f"Added keywords to skills: {skills_modified}")
        
        # Step 2: Find a few relevant bullet points and add context
        bullets_modified = enhance_relevant_bullets(doc, keywords_to_add)
        if bullets_modified:
            modifications.append(f"Enhanced {bullets_modified} bullet points")
        
        print(f"Modifications: {modifications}")
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
        
    except Exception as e:
        print(f"Error in docx editor: {e}")
        import traceback
        traceback.print_exc()
        # Return original if there's any error
        return original_docx_bytes


def append_to_skills_section(doc, keywords: List[str]) -> List[str]:
    """
    Find the skills section and append keywords that aren't already there.
    Returns list of keywords that were added.
    """
    skills_headers = ['skills', 'technical skills', 'technical expertise', 
                     'core competencies', 'expertise']
    
    added_keywords = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        
        # Look for skills header
        is_skills_header = any(h in text for h in skills_headers) and len(para.text) < 50
        
        if is_skills_header:
            print(f"Found skills header at paragraph {i}: '{para.text[:50]}'")
            
            # Look at the next few paragraphs for skills content
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                content_para = doc.paragraphs[j]
                content = content_para.text.strip()
                
                # Skip empty paragraphs
                if len(content) < 5:
                    continue
                
                # Stop if we hit another section header
                if is_section_header(content):
                    break
                
                # This is a skills line - check what we can add
                content_lower = content.lower()
                
                # Find keywords not already present
                new_keywords = []
                for kw in keywords:
                    if kw.lower() not in content_lower:
                        new_keywords.append(kw)
                
                if new_keywords:
                    # Append to this paragraph (max 5 at a time to keep it readable)
                    keywords_to_add = new_keywords[:5]
                    
                    # Detect the separator used
                    if ',' in content:
                        sep = ', '
                    else:
                        sep = ', '
                    
                    # Get formatting from existing runs
                    if content_para.runs:
                        last_run = content_para.runs[-1]
                        # Add a new run with the same formatting
                        new_run = content_para.add_run(sep + sep.join(keywords_to_add))
                        
                        # Copy formatting from existing run
                        if last_run.font.name:
                            new_run.font.name = last_run.font.name
                        if last_run.font.size:
                            new_run.font.size = last_run.font.size
                        if last_run.font.bold is not None:
                            new_run.font.bold = last_run.font.bold
                    else:
                        content_para.add_run(sep + sep.join(keywords_to_add))
                    
                    added_keywords.extend(keywords_to_add)
                    print(f"  Added to paragraph {j}: {keywords_to_add}")
                    
                    # Remove the added keywords from the list so we can add remaining to next line
                    keywords = [k for k in keywords if k not in keywords_to_add]
                    
                    # If all keywords added, stop
                    if not keywords:
                        return added_keywords
                    
                    # Only modify one skills line to be safe
                    return added_keywords
            
            break
    
    return added_keywords


def enhance_relevant_bullets(doc, keywords: List[str]) -> int:
    """
    Find bullet points that could benefit from keywords and add them naturally.
    Only modifies a maximum of 2 bullets to avoid over-optimization.
    """
    experience_headers = ['experience', 'work experience', 'professional experience', 
                         'employment', 'work history']
    
    in_experience = False
    bullets_modified = 0
    max_modifications = 2
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_lower = text.lower()
        
        # Check if we're entering experience section
        if any(h in text_lower for h in experience_headers) and len(text) < 50:
            in_experience = True
            continue
        
        # Check if we're leaving experience section
        if in_experience and is_section_header(text) and not any(h in text_lower for h in experience_headers):
            in_experience = False
            continue
        
        # Look for bullet points in experience section
        if in_experience and (text.startswith('•') or text.startswith('-') or text.startswith('*')):
            if bullets_modified >= max_modifications:
                break
            
            # Check if any keyword is relevant to this bullet
            for kw in keywords:
                kw_lower = kw.lower()
                
                # Skip if keyword already present
                if kw_lower in text_lower:
                    continue
                
                # Check if the bullet is related to the keyword
                # (simple heuristic: they share some common words)
                bullet_words = set(text_lower.split())
                kw_words = set(kw_lower.split())
                
                # If there's some relevance, add the keyword
                related_terms = {
                    'design': ['designed', 'developed', 'created', 'built'],
                    'analog': ['circuit', 'amplifier', 'comparator', 'bandgap'],
                    'verification': ['verified', 'tested', 'validated', 'simulated'],
                    'power': ['voltage', 'current', 'supply', 'pump'],
                }
                
                is_relevant = False
                for term, related in related_terms.items():
                    if term in kw_lower:
                        if any(r in text_lower for r in related):
                            is_relevant = True
                            break
                
                if is_relevant:
                    # Add keyword naturally at the end
                    if para.runs:
                        last_run = para.runs[-1]
                        
                        # Add keyword phrase
                        addition = f" utilizing {kw}"
                        if text.endswith('.'):
                            # Remove the period from last run, add keyword, then period
                            if last_run.text.endswith('.'):
                                last_run.text = last_run.text[:-1]
                            addition = f" utilizing {kw}."
                        
                        new_run = para.add_run(addition)
                        
                        # Copy formatting
                        if last_run.font.name:
                            new_run.font.name = last_run.font.name
                        if last_run.font.size:
                            new_run.font.size = last_run.font.size
                        
                        bullets_modified += 1
                        print(f"  Enhanced bullet at paragraph {i} with '{kw}'")
                        break  # Only add one keyword per bullet
    
    return bullets_modified


def is_section_header(text: str) -> bool:
    """Check if text looks like a section header."""
    headers = ['experience', 'education', 'skills', 'projects', 'certifications',
               'publications', 'awards', 'references', 'summary', 'objective',
               'work history', 'employment', 'qualifications', 'achievements',
               'technical expertise', 'technical skills', 'core competencies']
    
    text_lower = text.lower().strip()
    
    # Short text containing header keyword
    if len(text) < 50:
        for h in headers:
            if h in text_lower:
                return True
    
    return False


def enhance_resume_with_keywords(
    original_docx_bytes: bytes,
    keywords: List[str],
    improvements: List[str]
) -> bytes:
    """Wrapper for backward compatibility."""
    return create_tailored_resume_from_template(
        original_docx_bytes,
        {'keywords': keywords}
    )
