"""
File Parser Service - Extracts text from PDF and DOCX files
"""

import io
from typing import Optional


def parse_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file
    """
    try:
        from pypdf import PdfReader
        
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_content: bytes) -> str:
    """
    Extract text from a Word document (.docx)
    """
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from a file based on its extension
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return parse_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return parse_docx(file_content)
    elif filename_lower.endswith('.doc'):
        raise ValueError("Old .doc format not supported. Please save as .docx")
    elif filename_lower.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please use PDF, DOCX, or TXT.")


if __name__ == "__main__":
    # Quick test
    print("File parser module loaded successfully")
